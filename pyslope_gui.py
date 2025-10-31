"""
PySlope GUI Application

A tkinter interface for running PySlope slope stability analysis.

"""

import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox, filedialog
import matplotlib
matplotlib.use('TkAgg')
from matplotlib.figure import Figure
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import io
import base64
from PIL import Image
import tempfile
import os
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.section import WD_SECTION_START
    from docx.shared import Pt, Mm, RGBColor
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Add pyslope to path to import from local installation
sys.path.insert(0, os.path.abspath('.'))

try:
    from pyslope import Slope, Material, Udl, LineLoad
except ImportError:
    print("Please install pyslope: pip install pyslope")
    raise


class PySlopeGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("PySlope - Slope Stability Analysis")
        self.root.geometry("1400x900")
        
        # Data storage
        self.materials = []
        self.udls = []
        self.line_loads = []
        self.slope_obj = None
        self.current_plot = None
        
        # Create main notebook for tabs
        self.notebook = ttk.Notebook(root)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Create tabs
        self.create_project_tab()
        self.create_slope_tab()
        self.create_materials_tab()
        self.create_loads_tab()
        self.create_analysis_tab()
        self.create_results_tab()
        
        # Load default example values
        self.load_default_example()
    
    def create_project_tab(self):
        """Create tab for project details"""
        project_frame = ttk.Frame(self.notebook)
        self.notebook.add(project_frame, text="Project Details")
        
        # Project information
        info_frame = ttk.LabelFrame(project_frame, text="Project Information", padding=10)
        info_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        ttk.Label(info_frame, text="Project Name:").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.project_name_var = tk.StringVar(value="Slope Stability Analysis")
        ttk.Entry(info_frame, textvariable=self.project_name_var, width=40).grid(row=0, column=1, pady=5, padx=5)
        
        ttk.Label(info_frame, text="Project Reference:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.project_ref_var = tk.StringVar(value="25000")
        ttk.Entry(info_frame, textvariable=self.project_ref_var, width=40).grid(row=1, column=1, pady=5, padx=5)
        
        ttk.Label(info_frame, text="Project Location:").grid(row=2, column=0, sticky=tk.W, pady=5)
        self.project_location_var = tk.StringVar(value="8 Galah Grove, Nambucca Heads")
        ttk.Entry(info_frame, textvariable=self.project_location_var, width=40).grid(row=2, column=1, pady=5, padx=5)
        
        # Client information
        client_frame = ttk.LabelFrame(project_frame, text="Client Information", padding=10)
        client_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(client_frame, text="Client Name:").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.client_name_var = tk.StringVar(value="Tim Polo")
        ttk.Entry(client_frame, textvariable=self.client_name_var, width=40).grid(row=0, column=1, pady=5, padx=5)
        
        ttk.Label(client_frame, text="Client Company:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.client_company_var = tk.StringVar(value="")
        ttk.Entry(client_frame, textvariable=self.client_company_var, width=40).grid(row=1, column=1, pady=5, padx=5)
        
        ttk.Label(client_frame, text="Client Address:").grid(row=2, column=0, sticky=tk.W, pady=5)
        self.client_address_var = tk.StringVar(value="3a Nyora Close, Coffs Harbour")
        ttk.Entry(client_frame, textvariable=self.client_address_var, width=40).grid(row=2, column=1, pady=5, padx=5)
        
        # Engineer information
        engineer_frame = ttk.LabelFrame(project_frame, text="Engineer Information", padding=10)
        engineer_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(engineer_frame, text="Engineer Name:").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.engineer_name_var = tk.StringVar(value="Tim Polo")
        ttk.Entry(engineer_frame, textvariable=self.engineer_name_var, width=40).grid(row=0, column=1, pady=5, padx=5)
        
        ttk.Label(engineer_frame, text="Company:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.engineer_company_var = tk.StringVar(value="Nambucca Engineering")
        ttk.Entry(engineer_frame, textvariable=self.engineer_company_var, width=40).grid(row=1, column=1, pady=5, padx=5)
        
        ttk.Label(engineer_frame, text="Email:").grid(row=2, column=0, sticky=tk.W, pady=5)
        self.engineer_email_var = tk.StringVar(value="tim@nambuccaeng.com")
        ttk.Entry(engineer_frame, textvariable=self.engineer_email_var, width=40).grid(row=2, column=1, pady=5, padx=5)
        
        ttk.Label(engineer_frame, text="Phone:").grid(row=3, column=0, sticky=tk.W, pady=5)
        self.engineer_phone_var = tk.StringVar(value="0449646713")
        ttk.Entry(engineer_frame, textvariable=self.engineer_phone_var, width=40).grid(row=3, column=1, pady=5, padx=5)
        
    def create_slope_tab(self):
        """Create tab for slope definition"""
        slope_frame = ttk.Frame(self.notebook)
        self.notebook.add(slope_frame, text="Slope Definition")
        
        # Slope parameters
        params_frame = ttk.LabelFrame(slope_frame, text="Slope Parameters", padding=10)
        params_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        ttk.Label(params_frame, text="Height (m):").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.height_var = tk.StringVar(value="3")
        ttk.Entry(params_frame, textvariable=self.height_var, width=15).grid(row=0, column=1, pady=5, padx=5)
        
        ttk.Label(params_frame, text="Angle (deg):").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.angle_var = tk.StringVar(value="30")
        ttk.Entry(params_frame, textvariable=self.angle_var, width=15).grid(row=1, column=1, pady=5, padx=5)
        
        ttk.Label(params_frame, text="Length (m):").grid(row=2, column=0, sticky=tk.W, pady=5)
        self.length_var = tk.StringVar(value="")
        ttk.Entry(params_frame, textvariable=self.length_var, width=15).grid(row=2, column=1, pady=5, padx=5)
        ttk.Label(params_frame, text="(Leave empty if using angle)", font=("Arial", 8)).grid(row=2, column=2, sticky=tk.W, padx=5)
        
        # Uphill angle parameter
        ttk.Label(params_frame, text="Uphill Angle (deg):").grid(row=3, column=0, sticky=tk.W, pady=5)
        self.uphill_angle_var = tk.StringVar(value="")
        ttk.Entry(params_frame, textvariable=self.uphill_angle_var, width=15).grid(row=3, column=1, pady=5, padx=5)
        ttk.Label(params_frame, text="(Leave empty for flat, + = upward, - = downward)", font=("Arial", 8)).grid(row=3, column=2, sticky=tk.W, padx=5)
        
        # Water table
        water_frame = ttk.LabelFrame(slope_frame, text="Water Table", padding=10)
        water_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(water_frame, text="Depth from top of slope (m):").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.water_table_var = tk.StringVar(value="4")
        ttk.Entry(water_frame, textvariable=self.water_table_var, width=15).grid(row=0, column=1, pady=5, padx=5)
        ttk.Label(water_frame, text="(Leave empty for no water table)", font=("Arial", 8)).grid(row=0, column=2, sticky=tk.W, padx=5)
        
        # Analysis limits
        limits_frame = ttk.LabelFrame(slope_frame, text="Analysis Limits", padding=10)
        limits_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(limits_frame, text="Left limit (m):").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.left_limit_var = tk.StringVar(value="")
        ttk.Entry(limits_frame, textvariable=self.left_limit_var, width=15).grid(row=0, column=1, pady=5, padx=5)
        
        ttk.Label(limits_frame, text="Right limit (m):").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.right_limit_var = tk.StringVar(value="")
        ttk.Entry(limits_frame, textvariable=self.right_limit_var, width=15).grid(row=1, column=1, pady=5, padx=5)
        ttk.Label(limits_frame, text="(Leave empty to use defaults)", font=("Arial", 8)).grid(row=1, column=2, sticky=tk.W, padx=5)
        
        # Analysis options
        options_frame = ttk.LabelFrame(slope_frame, text="Analysis Options", padding=10)
        options_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(options_frame, text="Number of slices:").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.slices_var = tk.StringVar(value="50")
        ttk.Entry(options_frame, textvariable=self.slices_var, width=15).grid(row=0, column=1, pady=5, padx=5)
        
        ttk.Label(options_frame, text="Number of iterations:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.iterations_var = tk.StringVar(value="2000")
        ttk.Entry(options_frame, textvariable=self.iterations_var, width=15).grid(row=1, column=1, pady=5, padx=5)
        
    def create_materials_tab(self):
        """Create tab for materials management"""
        materials_frame = ttk.Frame(self.notebook)
        self.notebook.add(materials_frame, text="Materials")
        
        # Add material form
        form_frame = ttk.LabelFrame(materials_frame, text="Add Material", padding=10)
        form_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(form_frame, text="Unit Weight (kN/m³):").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.mat_unit_weight_var = tk.StringVar()
        ttk.Entry(form_frame, textvariable=self.mat_unit_weight_var, width=15).grid(row=0, column=1, pady=5, padx=5)
        
        ttk.Label(form_frame, text="Friction Angle (deg):").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.mat_friction_var = tk.StringVar()
        ttk.Entry(form_frame, textvariable=self.mat_friction_var, width=15).grid(row=1, column=1, pady=5, padx=5)
        
        ttk.Label(form_frame, text="Cohesion (kPa):").grid(row=2, column=0, sticky=tk.W, pady=5)
        self.mat_cohesion_var = tk.StringVar()
        ttk.Entry(form_frame, textvariable=self.mat_cohesion_var, width=15).grid(row=2, column=1, pady=5, padx=5)
        
        ttk.Label(form_frame, text="Depth to Bottom (m):").grid(row=3, column=0, sticky=tk.W, pady=5)
        self.mat_depth_var = tk.StringVar()
        ttk.Entry(form_frame, textvariable=self.mat_depth_var, width=15).grid(row=3, column=1, pady=5, padx=5)
        
        ttk.Button(form_frame, text="Add Material", command=self.add_material).grid(row=4, column=0, columnspan=2, pady=10)
        
        # Materials list
        list_frame = ttk.LabelFrame(materials_frame, text="Materials List", padding=10)
        list_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Treeview for materials
        columns = ("Unit Weight", "Friction Angle", "Cohesion", "Depth to Bottom")
        self.materials_tree = ttk.Treeview(list_frame, columns=columns, show="headings", height=10)
        for col in columns:
            self.materials_tree.heading(col, text=col)
            self.materials_tree.column(col, width=120)
        self.materials_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, command=self.materials_tree.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.materials_tree.configure(yscrollcommand=scrollbar.set)
        
        ttk.Button(list_frame, text="Remove Selected", command=self.remove_material).pack(pady=5)
        
    def create_loads_tab(self):
        """Create tab for loads (UDLs and Line Loads)"""
        loads_frame = ttk.Frame(self.notebook)
        self.notebook.add(loads_frame, text="Loads")
        
        # UDL section
        udl_frame = ttk.LabelFrame(loads_frame, text="Uniform Distributed Loads (UDL)", padding=10)
        udl_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(udl_frame, text="Magnitude (kPa):").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.udl_magnitude_var = tk.StringVar()
        ttk.Entry(udl_frame, textvariable=self.udl_magnitude_var, width=15).grid(row=0, column=1, pady=5, padx=5)
        
        ttk.Label(udl_frame, text="Offset from crest (m):").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.udl_offset_var = tk.StringVar(value="0")
        ttk.Entry(udl_frame, textvariable=self.udl_offset_var, width=15).grid(row=1, column=1, pady=5, padx=5)
        
        ttk.Label(udl_frame, text="Length (m):").grid(row=2, column=0, sticky=tk.W, pady=5)
        self.udl_length_var = tk.StringVar()
        ttk.Entry(udl_frame, textvariable=self.udl_length_var, width=15).grid(row=2, column=1, pady=5, padx=5)
        ttk.Label(udl_frame, text="(Leave empty for infinite)", font=("Arial", 8)).grid(row=2, column=2, sticky=tk.W, padx=5)
        
        ttk.Button(udl_frame, text="Add UDL", command=self.add_udl).grid(row=3, column=0, columnspan=2, pady=10)
        
        # UDL list
        udl_list_frame = ttk.LabelFrame(loads_frame, text="UDL List", padding=10)
        udl_list_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        udl_columns = ("Magnitude", "Offset", "Length")
        self.udl_tree = ttk.Treeview(udl_list_frame, columns=udl_columns, show="headings", height=5)
        for col in udl_columns:
            self.udl_tree.heading(col, text=col)
            self.udl_tree.column(col, width=150)
        self.udl_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        udl_scrollbar = ttk.Scrollbar(udl_list_frame, orient=tk.VERTICAL, command=self.udl_tree.yview)
        udl_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.udl_tree.configure(yscrollcommand=udl_scrollbar.set)
        
        ttk.Button(udl_list_frame, text="Remove Selected", command=self.remove_udl).pack(pady=5)
        
        # Line Load section
        ll_frame = ttk.LabelFrame(loads_frame, text="Line Loads", padding=10)
        ll_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(ll_frame, text="Magnitude (kN/m):").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.ll_magnitude_var = tk.StringVar()
        ttk.Entry(ll_frame, textvariable=self.ll_magnitude_var, width=15).grid(row=0, column=1, pady=5, padx=5)
        
        ttk.Label(ll_frame, text="Offset from crest (m):").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.ll_offset_var = tk.StringVar(value="0")
        ttk.Entry(ll_frame, textvariable=self.ll_offset_var, width=15).grid(row=1, column=1, pady=5, padx=5)
        
        ttk.Button(ll_frame, text="Add Line Load", command=self.add_line_load).grid(row=2, column=0, columnspan=2, pady=10)
        
        # Line Load list
        ll_list_frame = ttk.LabelFrame(loads_frame, text="Line Load List", padding=10)
        ll_list_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        ll_columns = ("Magnitude", "Offset")
        self.ll_tree = ttk.Treeview(ll_list_frame, columns=ll_columns, show="headings", height=5)
        for col in ll_columns:
            self.ll_tree.heading(col, text=col)
            self.ll_tree.column(col, width=200)
        self.ll_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        ll_scrollbar = ttk.Scrollbar(ll_list_frame, orient=tk.VERTICAL, command=self.ll_tree.yview)
        ll_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.ll_tree.configure(yscrollcommand=ll_scrollbar.set)
        
        ttk.Button(ll_list_frame, text="Remove Selected", command=self.remove_line_load).pack(pady=5)
        
    def create_analysis_tab(self):
        """Create tab for running analysis"""
        analysis_frame = ttk.Frame(self.notebook)
        self.notebook.add(analysis_frame, text="Run Analysis")
        
        # Run button
        button_frame = ttk.Frame(analysis_frame)
        button_frame.pack(fill=tk.X, padx=10, pady=20)
        
        ttk.Button(button_frame, text="Run Slope Stability Analysis", 
                  command=self.run_analysis, style="Accent.TButton").pack(pady=10)
        
        # Analysis options on this tab too
        options_frame = ttk.LabelFrame(analysis_frame, text="Quick Options Override", padding=10)
        options_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(options_frame, text="Plot Type:").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.plot_type_var = tk.StringVar(value="critical")
        plot_combo = ttk.Combobox(options_frame, textvariable=self.plot_type_var, 
                                  values=["boundary", "critical", "all_planes"], 
                                  state="readonly", width=20)
        plot_combo.grid(row=0, column=1, pady=5, padx=5)
        
        ttk.Label(options_frame, text="Max FOS (for all planes plot):").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.max_fos_var = tk.StringVar(value="2.0")
        ttk.Entry(options_frame, textvariable=self.max_fos_var, width=15).grid(row=1, column=1, pady=5, padx=5)
        
    def create_results_tab(self):
        """Create tab for displaying results"""
        results_frame = ttk.Frame(self.notebook)
        self.notebook.add(results_frame, text="Results")
        
        # Results text
        text_frame = ttk.LabelFrame(results_frame, text="Analysis Results", padding=10)
        text_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        self.results_text = scrolledtext.ScrolledText(text_frame, height=10, wrap=tk.WORD)
        self.results_text.pack(fill=tk.BOTH, expand=True)
        
        # Plot display - store reference
        plot_controls = ttk.Frame(results_frame)
        plot_controls.pack(fill=tk.X, padx=10, pady=5)
        
        ttk.Button(plot_controls, text="Update Plot", command=self.update_plot).pack(side=tk.LEFT, padx=5)
        ttk.Button(plot_controls, text="Generate Report", command=self.generate_report).pack(side=tk.LEFT, padx=5)
        
        self.plot_frame = ttk.LabelFrame(results_frame, text="Plot", padding=10)
        self.plot_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        self.plot_canvas = None
        self.plot_label = None
        
    def add_material(self):
        """Add a material to the list"""
        try:
            unit_weight = float(self.mat_unit_weight_var.get())
            friction_angle = float(self.mat_friction_var.get())
            cohesion = float(self.mat_cohesion_var.get())
            depth_to_bottom = float(self.mat_depth_var.get())
            
            material_data = {
                'unit_weight': unit_weight,
                'friction_angle': friction_angle,
                'cohesion': cohesion,
                'depth_to_bottom': depth_to_bottom
            }
            
            self.materials.append(material_data)
            self.materials_tree.insert("", tk.END, values=(
                f"{unit_weight:.2f}", f"{friction_angle:.2f}", 
                f"{cohesion:.2f}", f"{depth_to_bottom:.2f}"
            ))
            
            # Clear form
            self.mat_unit_weight_var.set("")
            self.mat_friction_var.set("")
            self.mat_cohesion_var.set("")
            self.mat_depth_var.set("")
            
        except ValueError:
            messagebox.showerror("Error", "Please enter valid numeric values for all material properties.")
    
    def remove_material(self):
        """Remove selected material from list"""
        selected = self.materials_tree.selection()
        if selected:
            item = selected[0]
            index = self.materials_tree.index(item)
            del self.materials[index]
            self.materials_tree.delete(item)
    
    def add_udl(self):
        """Add a UDL to the list"""
        try:
            magnitude = float(self.udl_magnitude_var.get())
            offset = float(self.udl_offset_var.get()) if self.udl_offset_var.get() else 0
            length_str = self.udl_length_var.get()
            length = float(length_str) if length_str else None
            
            udl_data = {
                'magnitude': magnitude,
                'offset': offset,
                'length': length
            }
            
            self.udls.append(udl_data)
            length_display = f"{length:.2f}" if length else "Infinite"
            self.udl_tree.insert("", tk.END, values=(
                f"{magnitude:.2f}", f"{offset:.2f}", length_display
            ))
            
            # Clear form
            self.udl_magnitude_var.set("")
            self.udl_offset_var.set("0")
            self.udl_length_var.set("")
            
        except ValueError:
            messagebox.showerror("Error", "Please enter valid numeric values.")
    
    def remove_udl(self):
        """Remove selected UDL from list"""
        selected = self.udl_tree.selection()
        if selected:
            item = selected[0]
            index = self.udl_tree.index(item)
            del self.udls[index]
            self.udl_tree.delete(item)
    
    def add_line_load(self):
        """Add a line load to the list"""
        try:
            magnitude = float(self.ll_magnitude_var.get())
            offset = float(self.ll_offset_var.get()) if self.ll_offset_var.get() else 0
            
            ll_data = {
                'magnitude': magnitude,
                'offset': offset
            }
            
            self.line_loads.append(ll_data)
            self.ll_tree.insert("", tk.END, values=(
                f"{magnitude:.2f}", f"{offset:.2f}"
            ))
            
            # Clear form
            self.ll_magnitude_var.set("")
            self.ll_offset_var.set("0")
            
        except ValueError:
            messagebox.showerror("Error", "Please enter valid numeric values.")
    
    def remove_line_load(self):
        """Remove selected line load from list"""
        selected = self.ll_tree.selection()
        if selected:
            item = selected[0]
            index = self.ll_tree.index(item)
            del self.line_loads[index]
            self.ll_tree.delete(item)
    
    def run_analysis(self):
        """Run the slope stability analysis"""
        try:
            # Clear previous results
            self.results_text.delete(1.0, tk.END)
            if self.plot_canvas:
                self.plot_canvas.get_tk_widget().destroy()
            if self.plot_label:
                self.plot_label.destroy()
                self.plot_label = None
            
            # Get slope parameters
            height = float(self.height_var.get())
            angle_str = self.angle_var.get()
            length_str = self.length_var.get()
            uphill_angle_str = self.uphill_angle_var.get()
            
            angle = float(angle_str) if angle_str else None
            length = float(length_str) if length_str else None
            uphill_angle = float(uphill_angle_str) if uphill_angle_str else None
            
            # Create slope
            self.slope_obj = Slope(height=height, angle=angle, length=length, uphill_angle=uphill_angle)
            
            # Add materials
            if not self.materials:
                messagebox.showwarning("Warning", "Please add at least one material.")
                return
            
            material_objects = []
            for mat in self.materials:
                m = Material(
                    unit_weight=mat['unit_weight'],
                    friction_angle=mat['friction_angle'],
                    cohesion=mat['cohesion'],
                    depth_to_bottom=mat['depth_to_bottom']
                )
                material_objects.append(m)
            
            self.slope_obj.set_materials(*material_objects)
            
            # Add UDLs
            if self.udls:
                udl_objects = []
                for udl in self.udls:
                    u = Udl(
                        magnitude=udl['magnitude'],
                        offset=udl['offset'],
                        length=udl['length']
                    )
                    udl_objects.append(u)
                self.slope_obj.set_udls(*udl_objects)
            
            # Add Line Loads
            if self.line_loads:
                ll_objects = []
                for ll in self.line_loads:
                    l = LineLoad(
                        magnitude=ll['magnitude'],
                        offset=ll['offset']
                    )
                    ll_objects.append(l)
                self.slope_obj.set_lls(*ll_objects)
            
            # Set water table
            water_table_str = self.water_table_var.get()
            if water_table_str:
                self.slope_obj.set_water_table(float(water_table_str))
            
            # Set analysis limits
            left_limit_str = self.left_limit_var.get()
            right_limit_str = self.right_limit_var.get()
            if left_limit_str and right_limit_str:
                self.slope_obj.set_analysis_limits(
                    float(left_limit_str), 
                    float(right_limit_str)
                )
            elif not left_limit_str and not right_limit_str:
                # Use defaults based on slope coordinates
                self.slope_obj.set_analysis_limits(
                    self.slope_obj.get_top_coordinates()[0] - 5,
                    self.slope_obj.get_bottom_coordinates()[0] + 5
                )
            else:
                messagebox.showwarning("Warning", "Both limits must be provided or both left empty for defaults.")
                return
            
            # Update analysis options
            slices = int(self.slices_var.get())
            iterations = int(self.iterations_var.get())
            self.slope_obj.update_analysis_options(slices=slices, iterations=iterations)
            
            # Run analysis
            self.results_text.insert(tk.END, "Running analysis...\n")
            self.root.update()
            
            self.slope_obj.analyse_slope()
            
            # Get results
            min_fos = self.slope_obj.get_min_FOS()
            
            # Display results
            self.results_text.insert(tk.END, "=" * 50 + "\n")
            self.results_text.insert(tk.END, "SLOPE STABILITY ANALYSIS RESULTS\n")
            self.results_text.insert(tk.END, "=" * 50 + "\n\n")
            self.results_text.insert(tk.END, f"Critical Factor of Safety: {min_fos:.4f}\n\n")
            self.results_text.insert(tk.END, f"Slope Height: {height} m\n")
            if angle:
                self.results_text.insert(tk.END, f"Slope Angle: {angle} deg\n")
            if length:
                self.results_text.insert(tk.END, f"Slope Length: {length} m\n")
            if uphill_angle is not None:
                self.results_text.insert(tk.END, f"Uphill Angle: {uphill_angle} deg\n")
            self.results_text.insert(tk.END, f"\nNumber of Materials: {len(self.materials)}\n")
            self.results_text.insert(tk.END, f"Number of UDLs: {len(self.udls)}\n")
            self.results_text.insert(tk.END, f"Number of Line Loads: {len(self.line_loads)}\n")
            if water_table_str:
                self.results_text.insert(tk.END, f"Water Table Depth: {water_table_str} m\n")
            self.results_text.insert(tk.END, f"\nAnalysis Options:\n")
            self.results_text.insert(tk.END, f"  Slices: {slices}\n")
            self.results_text.insert(tk.END, f"  Iterations: {iterations}\n")
            
            # Create plot
            self.create_plot()
            
            # Switch to results tab (index 5 now with project tab)
            self.notebook.select(5)
            
            messagebox.showinfo("Success", f"Analysis completed!\nCritical FOS: {min_fos:.4f}")
            
        except Exception as e:
            messagebox.showerror("Error", f"Analysis failed: {str(e)}")
            self.results_text.insert(tk.END, f"\nERROR: {str(e)}\n")
            import traceback
            self.results_text.insert(tk.END, f"\n{traceback.format_exc()}\n")
    
    def create_plot(self):
        """Create and display the plot"""
        try:
            plot_type = self.plot_type_var.get()
            fig = None
            
            if plot_type == "boundary":
                fig = self.slope_obj.plot_boundary()
            elif plot_type == "critical":
                fig = self.slope_obj.plot_critical()
            elif plot_type == "all_planes":
                # Check if analysis has been run and search results exist
                if not hasattr(self.slope_obj, '_search') or len(self.slope_obj._search) == 0:
                    messagebox.showwarning("Warning", "No analysis results found. Please run the analysis first.")
                    return
                
                max_fos = float(self.max_fos_var.get())
                
                # Debug: Check search results
                num_results = len(self.slope_obj._search)
                num_below_max = sum(1 for s in self.slope_obj._search if s.get("FOS") is not None and s.get("FOS", float('inf')) < max_fos)
                
                if num_results < 2:
                    messagebox.showwarning("Warning", 
                        f"Only {num_results} failure plane(s) found. The analysis may need more iterations or the search may be too restrictive.")
                elif num_below_max < 2:
                    messagebox.showwarning("Warning", 
                        f"Found {num_results} failure planes, but only {num_below_max} have FOS < {max_fos}. "
                        f"Consider increasing the max FOS value or checking the analysis results.")
                
                try:
                    fig = self.slope_obj.plot_all_planes(max_fos=max_fos)
                    # Verify that the figure has multiple traces (more than just the critical plane)
                    if fig and hasattr(fig, 'data') and len(fig.data) <= 2:
                        # Only boundary/material traces, no failure planes added
                        self.results_text.insert(tk.END, 
                            f"\nWarning: plot_all_planes only generated {len(fig.data)} traces. "
                            f"This may indicate that most failure planes were filtered out by max_fos={max_fos} "
                            f"or there was an issue generating the plot.\n")
                except IndexError as e:
                    messagebox.showerror("Error", 
                        f"Failed to generate all planes plot. This may indicate the analysis results are incomplete.\n"
                        f"Error: {str(e)}\n\nPlease try re-running the analysis.")
                    return
                except Exception as e:
                    messagebox.showerror("Error", 
                        f"Failed to generate all planes plot: {str(e)}\n\n"
                        f"Trying to generate critical plot instead.")
                    fig = self.slope_obj.plot_critical()
            
            if fig:
                # Convert plotly figure to image and display
                # Save to temporary file first
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                temp_file_path = temp_file.name
                temp_file.close()  # Close the file handle immediately
                
                # Write the image
                fig.write_image(temp_file_path)
                
                # Read and display image
                from PIL import Image, ImageTk
                img = Image.open(temp_file_path)
                
                # Copy image data into memory and close file
                import io
                img_bytes = io.BytesIO()
                img.save(img_bytes, format='PNG')
                img_bytes.seek(0)
                img.close()  # Close the file handle
                
                # Open from bytes instead of file
                img = Image.open(img_bytes)
                
                # Clear previous plot
                if self.plot_label:
                    self.plot_label.destroy()
                
                # Resize image to fit
                self.plot_frame.update_idletasks()
                frame_width = self.plot_frame.winfo_width() - 40
                frame_height = self.plot_frame.winfo_height() - 40
                
                # Ensure minimum dimensions
                if frame_width < 100:
                    frame_width = 800
                if frame_height < 100:
                    frame_height = 600
                
                # Maintain aspect ratio
                aspect_ratio = img.width / img.height
                if frame_width / frame_height > aspect_ratio:
                    display_height = frame_height
                    display_width = int(frame_height * aspect_ratio)
                else:
                    display_width = frame_width
                    display_height = int(frame_width / aspect_ratio)
                
                img_resized = img.resize((display_width, display_height), Image.Resampling.LANCZOS)
                photo = ImageTk.PhotoImage(img_resized)
                
                self.plot_label = tk.Label(self.plot_frame, image=photo)
                self.plot_label.image = photo  # Keep a reference
                self.plot_label.pack()
                
                # Clean up temp file - try with a small delay for Windows
                try:
                    import time
                    time.sleep(0.1)  # Small delay to ensure file handles are released
                    if os.path.exists(temp_file_path):
                        os.unlink(temp_file_path)
                except (OSError, PermissionError) as e:
                    # If we can't delete it immediately, try to delete it later
                    # Store path for cleanup on next run if needed
                    pass
                
        except Exception as e:
            messagebox.showerror("Error", f"Failed to create plot: {str(e)}")
            self.results_text.insert(tk.END, f"\nPlot Error: {str(e)}\n")
            import traceback
            self.results_text.insert(tk.END, f"\n{traceback.format_exc()}\n")
    
    def update_plot(self):
        """Update the plot without re-running analysis"""
        if self.slope_obj:
            self.create_plot()
        else:
            messagebox.showwarning("Warning", "Please run an analysis first.")
    
    def generate_report(self):
        """Generate a Word document report summarizing all inputs and outputs"""
        if not DOCX_AVAILABLE:
            messagebox.showerror("Error", "python-docx is not installed. Please install it with: pip install python-docx")
            return
        
        if not self.slope_obj:
            messagebox.showwarning("Warning", "Please run an analysis first before generating a report.")
            return
        
        # Ask user for save location
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")],
            initialfile=f"Slope_Stability_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        )
        
        if not file_path:
            return  # User cancelled
        
        try:
            self.create_report_document(file_path)
            messagebox.showinfo("Success", f"Report generated successfully!\nSaved to: {file_path}")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate report: {str(e)}")
            import traceback
            print(traceback.format_exc())
    
    def create_report_document(self, file_path):
        """Create the Word document report with all inputs and outputs"""
        doc = Document()
        
        # Set up document styling
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0, 0, 0)
        
        # Set page margins
        section = doc.sections[0]
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(20)
        section.right_margin = Mm(20)
        
        # Add header with logo
        header = section.header
        header_distance = Mm(5)
        section.header_distance = header_distance
        header_paragraph = header.paragraphs[0]
        header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Add logo if it exists
        logo_path = os.path.join(os.path.dirname(__file__), "assets", "nambuccaLogo.png")
        if os.path.exists(logo_path):
            try:
                run = header_paragraph.add_run()
                run.add_picture(logo_path, width=Mm(40))
            except Exception as e:
                print(f"Warning: Could not add logo: {e}")
        
        # Get current date
        current_date = datetime.now()
        day = current_date.day
        month = current_date.strftime("%B")
        year = current_date.year
        
        # Header with project reference
        project_ref = self.project_ref_var.get() or "N/A"
        header_text = f"Ref: {project_ref}"
        if self.client_name_var.get():
            header_text += f"\n{self.client_name_var.get()}"
        if self.client_company_var.get():
            header_text += f"\n{self.client_company_var.get()}"
        if self.client_address_var.get():
            header_text += f"\n{self.client_address_var.get()}"
        header_text += f"\n{day} {month} {year}"
        
        doc.add_paragraph(header_text)
        
        # Title
        project_name = self.project_name_var.get() or "Slope Stability Analysis"
        title = doc.add_heading(level=1)
        title_run = title.add_run(f"\n{project_name}")
        if self.project_location_var.get():
            title_run.add_text(f"\n{self.project_location_var.get()}")
        title_run.bold = True
        title_run.font.size = Pt(14)
        title_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Introduction paragraph
        doc.add_paragraph(
            f"\nThis report presents the results of a slope stability analysis conducted using Bishop's method "
            f"of slices. The analysis was performed to determine the factor of safety against slope failure."
        )
        
        # Project Details Section
        details_heading = doc.add_paragraph("Project Details:", style="Heading 2")
        for run in details_heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.underline = True
        
        details_para = doc.add_paragraph()
        details_para.add_run(f"Project Name: ").bold = True
        details_para.add_run(f"{project_name}\n")
        if project_ref != "N/A":
            details_para.add_run(f"Project Reference: ").bold = True
            details_para.add_run(f"{project_ref}\n")
        if self.project_location_var.get():
            details_para.add_run(f"Location: ").bold = True
            details_para.add_run(f"{self.project_location_var.get()}\n")
        if self.client_name_var.get():
            details_para.add_run(f"Client: ").bold = True
            details_para.add_run(f"{self.client_name_var.get()}\n")
        if self.client_company_var.get():
            details_para.add_run(f"Client Company: ").bold = True
            details_para.add_run(f"{self.client_company_var.get()}\n")
        
        # Slope Geometry Section
        geometry_heading = doc.add_paragraph("Slope Geometry:", style="Heading 2")
        for run in geometry_heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.underline = True
        
        geometry_para = doc.add_paragraph()
        height = float(self.height_var.get())
        angle_str = self.angle_var.get()
        length_str = self.length_var.get()
        uphill_angle_str = self.uphill_angle_var.get()
        
        geometry_para.add_run(f"Slope Height: ").bold = True
        geometry_para.add_run(f"{height} m\n")
        if angle_str:
            geometry_para.add_run(f"Slope Angle: ").bold = True
            geometry_para.add_run(f"{angle_str} degrees\n")
        if length_str:
            geometry_para.add_run(f"Slope Length: ").bold = True
            geometry_para.add_run(f"{length_str} m\n")
        if uphill_angle_str:
            geometry_para.add_run(f"Uphill Surface Angle: ").bold = True
            geometry_para.add_run(f"{uphill_angle_str} degrees\n")
        else:
            geometry_para.add_run(f"Uphill Surface: ").bold = True
            geometry_para.add_run(f"Flat\n")
        
        # Materials Section
        materials_heading = doc.add_paragraph("Material Properties:", style="Heading 2")
        for run in materials_heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.underline = True
        
        # Create materials table
        materials_table = doc.add_table(rows=1, cols=5)
        materials_table.style = 'Light Grid Accent 1'
        
        # Header row
        hdr_cells = materials_table.rows[0].cells
        hdr_cells[0].text = 'Layer'
        hdr_cells[1].text = 'Unit Weight (kN/m³)'
        hdr_cells[2].text = 'Friction Angle (deg)'
        hdr_cells[3].text = 'Cohesion (kPa)'
        hdr_cells[4].text = 'Depth to Bottom (m)'
        
        # Make header bold
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Add material rows
        for i, mat in enumerate(self.materials, 1):
            row_cells = materials_table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = f"{mat['unit_weight']:.2f}"
            row_cells[2].text = f"{mat['friction_angle']:.2f}"
            row_cells[3].text = f"{mat['cohesion']:.2f}"
            row_cells[4].text = f"{mat['depth_to_bottom']:.2f}"
        
        # Loading Conditions Section
        loads_heading = doc.add_paragraph("Loading Conditions:", style="Heading 2")
        for run in loads_heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.underline = True
        
        loads_para = doc.add_paragraph()
        if self.udls:
            loads_para.add_run("Uniform Distributed Loads (UDL):\n").bold = True
            for i, udl in enumerate(self.udls, 1):
                loads_para.add_run(f"  UDL {i}: ").bold = True
                length_display = f"{udl['length']:.2f} m" if udl['length'] else "Infinite"
                loads_para.add_run(f"Magnitude = {udl['magnitude']:.2f} kPa, Offset = {udl['offset']:.2f} m, Length = {length_display}\n")
        
        if self.line_loads:
            loads_para.add_run("\nLine Loads:\n").bold = True
            for i, ll in enumerate(self.line_loads, 1):
                loads_para.add_run(f"  Line Load {i}: ").bold = True
                loads_para.add_run(f"Magnitude = {ll['magnitude']:.2f} kN/m, Offset = {ll['offset']:.2f} m\n")
        
        if not self.udls and not self.line_loads:
            loads_para.add_run("No surface loads applied.\n")
        
        # Water Table Section
        water_heading = doc.add_paragraph("Water Table:", style="Heading 2")
        for run in water_heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.underline = True
        
        water_para = doc.add_paragraph()
        water_table_str = self.water_table_var.get()
        if water_table_str:
            water_para.add_run(f"Water table depth from top of slope: {water_table_str} m\n")
        else:
            water_para.add_run("No water table considered in the analysis.\n")
        
        # Analysis Parameters Section
        analysis_heading = doc.add_paragraph("Analysis Parameters:", style="Heading 2")
        for run in analysis_heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.underline = True
        
        analysis_para = doc.add_paragraph()
        slices = int(self.slices_var.get())
        iterations = int(self.iterations_var.get())
        analysis_para.add_run(f"Number of slices: ").bold = True
        analysis_para.add_run(f"{slices}\n")
        analysis_para.add_run(f"Number of iterations: ").bold = True
        analysis_para.add_run(f"{iterations}\n")
        analysis_para.add_run(f"Analysis method: ").bold = True
        analysis_para.add_run(f"Bishop's Method of Slices\n")
        
        # Analysis Limits
        left_limit_str = self.left_limit_var.get()
        right_limit_str = self.right_limit_var.get()
        if left_limit_str and right_limit_str:
            analysis_para.add_run(f"Analysis limits: ").bold = True
            analysis_para.add_run(f"Left = {left_limit_str} m, Right = {right_limit_str} m\n")
        else:
            analysis_para.add_run(f"Analysis limits: ").bold = True
            top_x = self.slope_obj.get_top_coordinates()[0]
            bot_x = self.slope_obj.get_bottom_coordinates()[0]
            analysis_para.add_run(f"Default (Left = {top_x - 5:.2f} m, Right = {bot_x + 5:.2f} m)\n")
        
        # Results Section
        results_heading = doc.add_paragraph("Analysis Results:", style="Heading 2")
        for run in results_heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.underline = True
        
        results_para = doc.add_paragraph()
        min_fos = self.slope_obj.get_min_FOS()
        circle_info = self.slope_obj.get_min_FOS_circle()
        endpoints = self.slope_obj.get_min_FOS_end_points()
        
        results_para.add_run(f"Critical Factor of Safety (FOS): ").bold = True
        results_para.add_run(f"{min_fos:.4f}\n\n")
        
        results_para.add_run("Critical Failure Surface Properties:\n").bold = True
        results_para.add_run(f"  Circle Centre: ({circle_info[0]:.3f}, {circle_info[1]:.3f}) m\n")
        results_para.add_run(f"  Circle Radius: {circle_info[2]:.3f} m\n")
        results_para.add_run(f"  Entry Point: ({endpoints[0][0]:.3f}, {endpoints[0][1]:.3f}) m\n")
        results_para.add_run(f"  Exit Point: ({endpoints[1][0]:.3f}, {endpoints[1][1]:.3f}) m\n")
        
        # Add figure from plot
        fig_heading = doc.add_paragraph("Figure 1: Critical Failure Surface", style="Heading 3")
        for run in fig_heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Generate and save the plot
        try:
            # Always use critical plot for the report
            fig = self.slope_obj.plot_critical()
            
            if fig:
                # Save plot to temporary file
                temp_plot_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                temp_plot_path = temp_plot_file.name
                temp_plot_file.close()
                
                # Write the image
                fig.write_image(temp_plot_path, width=1600, height=1000, scale=1)
                
                # Add image to document
                image_para = doc.add_paragraph()
                image_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = image_para.add_run()
                
                # Add the image with appropriate width (fits within page margins)
                # A4 width is 210mm, margins are 20mm each side, so max width is ~170mm
                run.add_picture(temp_plot_path, width=Mm(150))
                
                # Clean up temp file
                try:
                    if os.path.exists(temp_plot_path):
                        os.unlink(temp_plot_path)
                except Exception as e:
                    print(f"Warning: Could not delete temp plot file: {e}")
                    
        except Exception as e:
            # If plot generation fails, add a note instead
            error_para = doc.add_paragraph()
            error_para.add_run(f"Note: Could not generate figure. Error: {str(e)}\n")
            print(f"Warning: Could not add figure to report: {e}")
        
        # Interpretation
        interpretation_heading = doc.add_paragraph("Interpretation:", style="Heading 2")
        for run in interpretation_heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.underline = True
        
        interpretation_para = doc.add_paragraph()
        if min_fos < 1.0:
            interpretation_para.add_run(f"A factor of safety of {min_fos:.4f} indicates that the slope is unstable and failure is likely to occur.\n")
        elif min_fos < 1.3:
            interpretation_para.add_run(f"A factor of safety of {min_fos:.4f} indicates marginal stability. The slope may be at risk of failure.\n")
        elif min_fos < 1.5:
            interpretation_para.add_run(f"A factor of safety of {min_fos:.4f} indicates acceptable stability for temporary conditions.\n")
        else:
            interpretation_para.add_run(f"A factor of safety of {min_fos:.4f} indicates acceptable stability for permanent conditions.\n")
        
        interpretation_para.add_run(
            "\nNote: The analysis assumes circular failure surfaces and uses Bishop's simplified method. "
            "Results should be interpreted by a qualified geotechnical engineer in the context of site-specific conditions."
        )
        
        # Engineer Information
        engineer_heading = doc.add_paragraph("Engineer Information:", style="Heading 2")
        for run in engineer_heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.underline = True
        
        engineer_para = doc.add_paragraph()
        if self.engineer_name_var.get():
            engineer_para.add_run(f"Engineer: ").bold = True
            engineer_para.add_run(f"{self.engineer_name_var.get()}\n")
        if self.engineer_company_var.get():
            engineer_para.add_run(f"Company: ").bold = True
            engineer_para.add_run(f"{self.engineer_company_var.get()}\n")
        if self.engineer_email_var.get():
            engineer_para.add_run(f"Email: ").bold = True
            engineer_para.add_run(f"{self.engineer_email_var.get()}\n")
        if self.engineer_phone_var.get():
            engineer_para.add_run(f"Phone: ").bold = True
            engineer_para.add_run(f"{self.engineer_phone_var.get()}\n")
        
        engineer_para.add_run(f"\nDate: {day} {month} {year}\n")
        
        # Add signature
        engineer_name = self.engineer_name_var.get()
        signature_para = doc.add_paragraph()
        
        # Add "Yours sincerely," if engineer name exists
        if engineer_name:
            signature_para.add_run("\nYours sincerely,\n\n\n")
            signature_para.add_run(f"{engineer_name}\n")
            
            # Add signature image if available
            signature_path = os.path.join(os.path.dirname(__file__), "assets", "TimSignature.png")
            if os.path.exists(signature_path):
                try:
                    signature_run = signature_para.add_run()
                    signature_run.add_picture(signature_path, width=Mm(40))
                except Exception as e:
                    print(f"Warning: Could not add signature: {e}")
        
        # Save document
        doc.save(file_path)
    
    def load_default_example(self):
        """Load default example values from the README example"""
        # Materials are already set correctly in the GUI (height=3, angle=30)
        # Water table is already set to "4"
        
        # Add Material 1: unit_weight=20, friction_angle=45, cohesion=2, depth_to_bottom=2
        mat1_data = {
            'unit_weight': 20,
            'friction_angle': 45,
            'cohesion': 2,
            'depth_to_bottom': 2
        }
        self.materials.append(mat1_data)
        self.materials_tree.insert("", tk.END, values=(
            f"{mat1_data['unit_weight']:.2f}", 
            f"{mat1_data['friction_angle']:.2f}", 
            f"{mat1_data['cohesion']:.2f}", 
            f"{mat1_data['depth_to_bottom']:.2f}"
        ))
        
        # Add Material 2: unit_weight=20, friction_angle=30, cohesion=2, depth_to_bottom=5
        mat2_data = {
            'unit_weight': 20,
            'friction_angle': 30,
            'cohesion': 2,
            'depth_to_bottom': 5
        }
        self.materials.append(mat2_data)
        self.materials_tree.insert("", tk.END, values=(
            f"{mat2_data['unit_weight']:.2f}", 
            f"{mat2_data['friction_angle']:.2f}", 
            f"{mat2_data['cohesion']:.2f}", 
            f"{mat2_data['depth_to_bottom']:.2f}"
        ))
        
        # Add UDL 1: magnitude=100, offset=2, length=1
        udl1_data = {
            'magnitude': 100,
            'offset': 2,
            'length': 1
        }
        self.udls.append(udl1_data)
        self.udl_tree.insert("", tk.END, values=(
            f"{udl1_data['magnitude']:.2f}", 
            f"{udl1_data['offset']:.2f}", 
            f"{udl1_data['length']:.2f}"
        ))
        
        # Add UDL 2: magnitude=20, offset=0, length=None (infinite)
        udl2_data = {
            'magnitude': 20,
            'offset': 0,
            'length': None
        }
        self.udls.append(udl2_data)
        self.udl_tree.insert("", tk.END, values=(
            f"{udl2_data['magnitude']:.2f}", 
            f"{udl2_data['offset']:.2f}", 
            "Infinite"
        ))
        
        # Add Line Load: magnitude=10, offset=3
        ll1_data = {
            'magnitude': 10,
            'offset': 3
        }
        self.line_loads.append(ll1_data)
        self.ll_tree.insert("", tk.END, values=(
            f"{ll1_data['magnitude']:.2f}", 
            f"{ll1_data['offset']:.2f}"
        ))


def main():
    root = tk.Tk()
    app = PySlopeGUI(root)
    root.mainloop()


if __name__ == "__main__":
    main()

